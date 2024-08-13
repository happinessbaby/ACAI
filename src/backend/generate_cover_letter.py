# Import the necessary modules
import os
import openai
from utils.openai_api import get_completion
from utils.basic_utils import read_txt, process_json, write_file
from utils.common_utils import (retrieve_or_create_resume_info, retrieve_or_create_job_posting_info,
                                 retrieve_from_db,  search_related_samples)
from datetime import date
from pathlib import Path
import json
from json import JSONDecodeError
from typing import List
from utils.langchain_utils import create_summary_chain, generate_multifunction_response, handle_tool_error
from utils.agent_tools import create_search_tools, create_sample_tools
from docxtpl import DocxTemplate	
from docx import Document
from datetime import datetime
from io import BytesIO
from utils.aws_manager import get_client
import tempfile
from dotenv import load_dotenv, find_dotenv
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain_core.tools import Tool, tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

_ = load_dotenv(find_dotenv()) # read local .env file

openai.api_key = os.environ["OPENAI_API_KEY"]
cover_letter_samples_path = os.environ["COVER_LETTER_SAMPLES_PATH"]
faiss_web_data = os.environ["FAISS_WEB_DATA_PATH"]
STORAGE = os.environ["STORAGE"]
# TODO: caching and serialization of llm
llm = ChatOpenAI(temperature=0.9)
embeddings = OpenAIEmbeddings()
delimiter = "####"
delimiter1 = "****"
delimiter2 = "'''"
delimiter3 = '---'
delimiter4 = '////'
delimiter5 = '~~~~'

# local_save_path = os.environ["CHAT_PATH"]
if STORAGE=="CLOUD":
    bucket_name = os.environ["BUCKET_NAME"]
    s3_save_path = os.environ["S3_CHAT_PATH"]
    s3 = get_client('s3')
else:
    bucket_name=None
    s3=None
      
def generate_basic_cover_letter(resume_dict={}, job_posting_dict={}, output_dir=None, ) -> None:
    
    """ Main function that generates the cover letter.
    
    Keyword Args:

      about_me(str)s

      resume_file(str): path to the resume file in txt format

      posting_path(str): path to job posting in txt format
     
    """
    
    document = Document()
    document.add_heading('Cover Letter', 0)
    print(")0000000000000000)))))))))")
    end_path = os.path.join(output_dir, "cover_letter.docx")
    # Get resume info
    resume_content = resume_dict["resume_content"]
    work_experience_level = resume_dict.get("work_experience_level", "")
    qualifications = job_posting_dict.get("qualifications", [])
    responsibilities = job_posting_dict.get("responsibilities", [])
    job_specification = ", ".join(qualifications) + "\n" + ", ".join(responsibilities)
    company_description = job_posting_dict.get("company_description", "")
    company = job_posting_dict.get("company", "")
    job = resume_dict.get("pusuit_jobs", "")
    highest_education_level = resume_dict["education"].get("degree", "")
    name = resume_dict["contact"].get("name", "")
    phone = resume_dict["contact"].get("phone", "")
    email = resume_dict["contact"].get("email", "")
    print("AAAAAAAAAAAAAAAAAAAAAAAaaa")
    # Get adviced from web data on personalized best practices
    # advice_query = f"""Best practices when writing a cover letter for applicant with {highest_education_level} and {work_experience_level} experience as a {job}"""
    # advices = retrieve_from_db(advice_query, vectorstore_path=faiss_web_data, vectorstore_type="faiss")
    # # Get sample comparisons
    # related_samples = search_related_samples(job, cover_letter_samples_path)
    # sample_tools, tool_names = create_sample_tools(related_samples, "cover_letter")
    # # Get resume's relevant and irrelevant info for job: few-shot learning works great here
    # query_relevancy = f""" You are an expert resume advisor. 
    
    #  Step 1: Determine the relevant and irrelevant information contained in the resume document delimited with {delimiter} characters.

    #   resume: {delimiter}{resume_content}{delimiter} \n

    #   Generate a list of irrelevant information that should not be included in the cover letter and a list of relevant information that should be included in the cover letter.
       
    #   Remember to use either job specification or general job description as your guideline. 

    #   job specification: {job_specification}

    #     Your answer should be detailed and only from the resume. Please also provide your reasoning too. 
        
    #     For example, your answer may look like this:

    #     Relevant information:

    #     1. Experience as a Big Data Engineer: using Spark and Python are transferable skills to using Panda in data analysis

    #     Irrelevant information:

    #     1. Education in Management of Human Resources is not directly related to skills required for a data analyst 

    #   Step 2:  Sample cover letters are provided in your tools. Research {str(tool_names)} and answer the following question: and answer the following question:

    #        Make a list of common features these cover letters share. 

    #     """
    # # tool = [search_relevancy_advice]
    # relevancy = generate_multifunction_response(query_relevancy, sample_tools)
    # Use an LLM to generate a cover letter that is specific to the resume file that is being read
    # Step wise instructions: https://learn.deeplearning.ai/chatgpt-building-system/lesson/5/chain-of-thought-reasoning
      # Step 1: You are given two lists of information delimited with characters. One is irrelevant to applying for {job} and the other is relevant. 

      #   Use them as a reference when determining what to include and what not to include in the cover letter. 
    
      #   information list: {relevancy}.  \n

      # Step 2: You are also given some expert advices. Keep them in mind when generating the cover letter.

      #   expert advices: {advices}
    template_string2 = """You are a professional cover letter writer. A Human client has asked you to generate a cover letter for them.
  
        The content you are to use as reference to create the cover letter is delimited with {delimiter} characters.

        Always use this as a context when writing the cover letter. Do not write out of context and do not make anything up. Anything that should be included but unavailable can be put in brackets. 

        content: {delimiter}{content}{delimiter}. \n


      Step 1: You're provided with some company informtion, job description, and/or job specification. 

        Use it to make the cover letter cater to the company values. 

        company information: {company_description}.  \n

        job specification: {job_specification} \n
    
      Step 2: Change all personal information of the cover letter to the following. Do not incude them if they are -1 or empty: 

        name: {name}. \

        email: {email}. \

        phone number: {phone}. \
        
        today's date: {date}. \
        
        company they are applying to: {company}. \

        job position they are applying for: {job}. \
    
      Step 3: Generate the cover letter using what you've learned in Step 1 through Step 4. Do not make stuff up. 
    
      Use the following format:
        Step 1:{delimiter4} <step 1 reasoning>
        Step 2:{delimiter4} <step 2 reasoning>
        Step 3:{delimiter4} <step 3 reasoning>
        Step 4:{delimiter4} <step 4 reasoning>
        Step 5:{delimiter4} <the cover letter you generate>

      Make sure to include {delimiter4} to separate every step.
    """  
    prompt_template = ChatPromptTemplate.from_template(template_string2)
    print("BBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBBB")
    # print(prompt_template.messages[0].prompt.input_variables)
    cover_letter_message = prompt_template.format_messages(
                    name = name,
                    phone = phone,
                    email = email,
                    date = date.today(),
                    company = company,
                    job = job,
                    content=resume_content,
                    # relevancy=relevancy, 
                    # advices = advices,
                    company_description = company_description, 
                    job_specification = job_specification,
                    delimiter = delimiter,
                    delimiter4 = delimiter4,
    )
    my_cover_letter = llm.invoke(cover_letter_message).content
    print("CCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCCC")
    cover_letter = get_completion(f"Extract the entire cover letter and nothing else in the following text: {my_cover_letter}")
    document.add_paragraph(cover_letter)
    print("DDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDD")
    if STORAGE=="LOCAL":
      document.save(end_path)
    elif STORAGE=="CLOUD":
          # Save the rendered template to a BytesIO object
        output_stream = BytesIO()
        document.save(output_stream)
        output_stream.seek(0)    
        print("EEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEEE")
        # # Upload the BytesIO object to S3
        # s3.put_object(Bucket=bucket_name, Key=end_path, Body=output_stream.getvalue())
        # Write to a temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file.write(output_stream.getvalue())
            temp_file.seek(0)  # Reset stream position to the beginning if needed
            end_path = temp_file.name
            print(f"Temporary file created at: {end_path}")
    return end_path


def generate_preformatted_cover_letter(resume_file, job_posting_file='', job_description='', save_path="./test_cover_letter.docx"):
    
    # Part 1: Insert user information into template
    template_file = "./backend/cover_letter_templates/template1.docx"
    cover_letter_template = DocxTemplate(template_file)
    resume_dict = retrieve_or_create_resume_info(resume_file) 
    job_posting_dict= retrieve_or_create_job_posting_info(posting_path=job_posting_file, about_job=job_description, )
    job_posting_dict.update(resume_dict["contact"])
    info_dict=job_posting_dict
    func = lambda key, default: default if key not in info_dict or info_dict[key]==-1 else info_dict[key]
    personal_context = {
        "NAME": func("name", "YOUR NAME"),
        "CITY": func("city", "YOUR CITY"),
        "STATE": func("state", "YOUR STATE"),
        "PHONE": func("phone", "YOUR PHONE"),
        "EMAIL": func("email", "YOUR EMAIL"),
        "LINKEDIN": func("linkedin", "YOUR LINKEDIN URL"),
        "WEBSITE": func("website", "YOUR WEBSITE"),
        "JOB": func("job", "JOB TITLE"),
        "COMPANY": func("company", "COMPANY"),
        "DATE": datetime.today().date()
    }
    cover_letter_template.render(personal_context)
    cover_letter_template.save(save_path) 

    # Part 2: Compose cover letter draft with additional information
    company_description=job_posting_dict.get("company description", "")
    duties = job_posting_dict.get("duties", "")
    traits=job_posting_dict.get("qualifications", "")
    important_keywords = job_posting_dict.get("frequent_words", "")
    soft_skills=job_posting_dict.get("soft_skills", "")
    hard_skills = job_posting_dict.get("hard_skills", "")
    resume_content = read_txt(resume_file, storage=STORAGE, bucket_name=bucket_name, s3=s3)
    relevant_hard_skills = research_relevancy_in_resume(resume_content, hard_skills, "hard skills")
    relevant_soft_skills = research_relevancy_in_resume(resume_content, soft_skills, "soft skills")
    relevant_responsibilities = research_relevancy_in_resume(resume_content, duties, "responsibilities")
    prompt = """You are a professional cover letter writer. A Human candidate has asked you to generate a cover letter for them using a template. The template is given below:
    
    Cover letter template {cover_letter_template}      

    There are still missing information in brackets that needs to be filled in. Your job is to help fill in the rest of the cover letter with the following information:

    Company mission: {company_mission} \
    
    candidate's hard skills qualification: {relevant_hard_skills}

    candidate's soft skills qualification: {relevant_soft_skills}

    candidate's responsiblities qualification: {relevant_responsibilities}
    
    Please output the cover letter only. 
    
    """
    tmp_filename= Path(save_path).stem+Path(save_path).suffix
    convert_doc_to_txt(save_path,  tmp_filename)
    cover_letter = read_txt(tmp_filename)

    prompt_template = ChatPromptTemplate.from_template(prompt)
    # print(prompt_template.messages[0].prompt.input_variables)
    cover_letter_message = prompt_template.format_messages(
                    cover_letter_template=cover_letter,
                    company_mission=company_description,
                    relevant_hard_skills=relevant_hard_skills,
                     relevant_soft_skills=relevant_soft_skills, 
                     relevant_responsibilities=relevant_responsibilities,)
    my_cover_letter = llm(cover_letter_message).content
    print(f"Sucessfully written cover letter: {my_cover_letter}")
    write_file(my_cover_letter, "./cl_final_test.txt", mode="w")

# @tool(return_direct=True)
# def cover_letter_generator(json_request:str) -> str:
    
#     """Helps to generate a cover letter. Use this tool more than any other tool when user asks to write a cover letter.  

#     Do not use this tool when user provides a cover letter file and asks you to customize or improve it. 

#     Input should be a single string strictly in the following JSON format: '{{"about me":"<about me>", "resume file":"<resume file>", "job post link":"<job post link>"}}' \n

#     Leave value blank if there's no information provided. DO NOT MAKE STUFF UP. 

#     (remember to respond with a markdown code snippet of a JSON blob with a single action, and NOTHING else) 

#     Output should be using the "get download link" tool in the following single string JSON format: '{{"file_path":"<file_path>"}}'

#     """
#     # Output should be the exact cover letter that's generated for the user word for word. 
    
#     try:
#       json_request = json_request.strip("'<>() ").replace('\'', '\"')
#       args = json.loads(json_request)
#     except JSONDecodeError as e:
#       print(f"JSON DECODE ERROR: {e}")
#       return "Format in a single string JSON and try again."
   
#     if ("resume_file" not in args or args["resume_file"]=="" or args["resume_file"]=="<resume_file>"):
#       return "Can you provide your resume so I can further assist you? "
#     else:
#         resume_file = args["resume_file"]
#     if ("about me" not in args or args["about_me"] == "" or args["about_me"]=="<about_me>"):
#         about_me = ""
#     else:
#         about_me = args["about_me"]
#     if ("job_post_file" not in args or args["job_post_file"]=="" or args["job_post_file"]=="<job_post_file>"):
#         posting_path = ""
#     else:
#         posting_path = args["job_post_file"]

#     return generate_basic_cover_letter(about_me=about_me, resume_file=resume_file, posting_path=posting_path)



def processing_cover_letter(json_request: str) -> None:
    
    """ Input parser: input is LLM's action_input in JSON format. This function then processes the JSON data and feeds them to the cover letter generator. """

    try:
      args = json.loads(process_json(json_request))
    except JSONDecodeError as e:
      print(f"JSON DECODE ERROR: {e}")
      return "Format in a single string JSON and try again."
    # if resume doesn't exist, ask for resume
    if ("resume_file" not in args or args["resume_file"]=="" or args["resume_file"]=="<resume_file>"):
      return "Stop using the cover letter generator tool. Ask user for their resume, along with any other additional information that they could provide. "
    else:
        resume_file = args["resume_file"]
    if ("about_me" not in args or args["about_me"] == "" or args["about_me"]=="<about_me>"):
        about_me = ""
    else:
        about_me = args["about_me"]
    if ("job_posting_file" not in args or args["job_posting_file"]=="" or args["job_posting_file"]=="<job_posting_file>"):
        posting_path = ""
    else:
        posting_path = args["job_posting_file"]
    

    return generate_basic_cover_letter(about_me=about_me, resume_file=resume_file, posting_path=posting_path)


    
def create_cover_letter_generator_tool() -> List[Tool]:
    
    """ Input parser: input is user's input as a string of text. This function takes in text and parses it into JSON format. 
    
    Then it calls the processing_cover_letter function to process the JSON data. """
    
    name = "cover_letter_generator"
    parameters = '{{"about_me":"<about_me>", "resume_file":"<resume_file>", "job_posting_file":"<job_posting_file>"}}'
    description = f"""Helps to generate a cover letter. Use this tool more than any other tool when user asks to write a cover letter. 
     Input should be a single string strictly in the following JSON format: {parameters} \n
    """
    tools = [
        Tool(
        name = name,
        func =processing_cover_letter,
        description = description, 
        verbose = True,
        handle_tool_error=handle_tool_error,
        )
    ]
    print("Sucessfully created cover letter generator tool. ")
    return tools

  
    
 
if __name__ == '__main__':
    # test run defaults, change for yours (only resume_file cannot be left empty)
    resume_file = "/home/tebblespc/GPT-Projects/ACAI/ACAI/src/my_material/resume2023v3.txt"
    posting_path= "/home/tebblespc/GPT-Projects/ACAI/ACAI/src/my_material/rov.txt"
    template_file = "/home/tebblespc/GPT-Projects/ACAI/ACAI/src/backend/resume_templates/functional/functional1.docx"
    generate_basic_cover_letter(resume_file = resume_file, posting_path=posting_path)



